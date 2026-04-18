"""
Script to read the UG Project Phase II template .doc file and print its content.
Run this script to extract the template structure.
"""
import sys
import os

doc_path = r"c:\Users\yogi8\AppData\Local\Packages\5319275A.WhatsAppDesktop_cv1g1gvanyjgm\LocalState\sessions\4F7B19B35DD117CFFA761EAC3CD54216F35AAF9A\transfers\2026-16\UG-ProjPh-II 26 Template -AI.doc"

# Try python-docx first
try:
    from docx import Document
    doc = Document(doc_path)
    print("=== TEMPLATE CONTENT (via python-docx) ===\n")
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"[Para {i}] Style={para.style.name!r}: {para.text}")
    
    # Also check tables
    for t_idx, table in enumerate(doc.tables):
        print(f"\n=== TABLE {t_idx} ===")
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    print(f"  CELL: {cell.text[:200]}")
    sys.exit(0)
except ImportError:
    print("python-docx not installed. Trying to install...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    doc = Document(doc_path)
    print("=== TEMPLATE CONTENT ===\n")
    for para in doc.paragraphs:
        if para.text.strip():
            print(f"[{para.style.name}] {para.text}")
    sys.exit(0)
except Exception as e:
    print(f"python-docx error: {e}")

# Fallback: read raw bytes and extract text
try:
    with open(doc_path, 'rb') as f:
        content = f.read()
    # Extract readable ASCII strings
    import re
    strings = re.findall(b'[\x20-\x7e]{4,}', content)
    print("=== RAW STRINGS FROM .doc ===")
    for s in strings:
        decoded = s.decode('ascii', errors='ignore').strip()
        if len(decoded) > 10:
            print(decoded)
except Exception as e2:
    print(f"Raw read error: {e2}")
